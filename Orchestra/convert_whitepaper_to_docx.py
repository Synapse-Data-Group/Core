"""
Convert Orchestra Benchmark Whitepaper from Markdown to DOCX with proper formatting.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_styled_document():
    """Create a document with custom styles."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Define custom styles
    styles = doc.styles
    
    # Title style
    if 'CustomTitle' not in styles:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Subtitle style
    if 'CustomSubtitle' not in styles:
        subtitle_style = styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.bold = True
        subtitle_style.font.color.rgb = RGBColor(64, 64, 64)
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(6)
    
    return doc

def parse_markdown_to_docx(md_file, docx_file):
    """Parse markdown file and convert to formatted DOCX."""
    doc = create_styled_document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
            else:
                # End of code block
                in_code_block = False
                if code_block_lines:
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table and not line.startswith('|'):
            # End of table
            in_table = False
            if table_lines:
                add_table_to_doc(doc, table_lines)
                table_lines = []
        
        # Skip horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle headers
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line.lstrip('#').strip()
            
            if level == 1:
                p = doc.add_paragraph(text, style='CustomTitle')
            elif level == 2:
                p = doc.add_paragraph(text)
                p.style = 'Heading 1'
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.bold = True
            elif level == 3:
                p = doc.add_paragraph(text)
                p.style = 'Heading 2'
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
            elif level == 4:
                p = doc.add_paragraph(text)
                p.style = 'Heading 3'
                p.runs[0].font.size = Pt(12)
                p.runs[0].font.bold = True
            else:
                p = doc.add_paragraph(text)
                p.runs[0].font.bold = True
            
            i += 1
            continue
        
        # Handle bold/italic text with subtitle detection
        if line.startswith('**') and line.endswith('**') and len(line) < 100:
            text = line.strip('*').strip()
            p = doc.add_paragraph(text, style='CustomSubtitle')
            i += 1
            continue
        
        # Handle italic text (single line)
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            text = line.strip('*').strip()
            p = doc.add_paragraph(text)
            p.runs[0].font.italic = True
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = format_inline_markdown(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = format_inline_markdown(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip():
            text = format_inline_markdown(line.strip())
            p = doc.add_paragraph()
            add_formatted_text(p, text)
        else:
            # Empty line
            doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ DOCX file created: {docx_file}")

def add_table_to_doc(doc, table_lines):
    """Add a markdown table to the document."""
    # Parse table
    rows = []
    for line in table_lines:
        if line.strip().startswith('|---') or line.strip().startswith('|-'):
            continue  # Skip separator line
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Populate table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            # Remove markdown formatting from cell
            cell_text = cell_data.replace('**', '').replace('*', '')
            cell.text = cell_text
            
            # Bold header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_paragraph()  # Add space after table

def format_inline_markdown(text):
    """Keep markdown formatting markers for later processing."""
    return text

def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, code)."""
    # Split by code blocks first
    parts = re.split(r'(`[^`]+`)', text)
    
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            # Code
            run = paragraph.add_run(part.strip('`'))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Process bold and italic
            subparts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', part)
            for subpart in subparts:
                if subpart.startswith('**') and subpart.endswith('**'):
                    # Bold
                    run = paragraph.add_run(subpart.strip('*'))
                    run.font.bold = True
                elif subpart.startswith('*') and subpart.endswith('*'):
                    # Italic
                    run = paragraph.add_run(subpart.strip('*'))
                    run.font.italic = True
                else:
                    # Regular text
                    if subpart:
                        paragraph.add_run(subpart)

if __name__ == '__main__':
    md_file = 'BENCHMARK_WHITEPAPER.md'
    docx_file = 'Orchestra_Benchmark_Whitepaper.docx'
    
    print(f"Converting {md_file} to {docx_file}...")
    parse_markdown_to_docx(md_file, docx_file)
    print("✅ Conversion complete!")
